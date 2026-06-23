"""Tests for CV document builder job-skill sections."""

import os
import tempfile
import unittest

from job_apply_ai.cv_modifier.docx_builder import CVDocumentBuilder
from job_apply_ai.storage.user_profile import get_default_cv_template_path


class DocxBuilderJobSkillsTests(unittest.TestCase):
    def test_export_omits_preview_only_job_skill_sections(self):
        template_path = get_default_cv_template_path()
        self.assertTrue(os.path.exists(template_path))

        content = {
            "professional_summary": "Experienced developer targeting this role.",
            "job_matched_skills": ["Python", "Android", "REST APIs"],
            "job_skills_not_in_cv": ["Kotlin", "GraphQL"],
            "technical_skills": ["Python", "Android", "Java", "Flask"],
            "tools_platforms": ["Git", "Docker"],
        }

        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = os.path.join(tmpdir, "tailored_cv.docx")
            builder = CVDocumentBuilder(template_path)
            builder.build(output_path, content)

            from docx import Document

            doc = Document(output_path)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

            self.assertNotIn("Skills Matching Job Description", paragraphs)
            self.assertNotIn("Job Skills Not In CV", paragraphs)

            technical_idx = paragraphs.index("Technical Skills")
            technical_line = paragraphs[technical_idx + 1]
            self.assertIn("Python", technical_line)
            self.assertIn("REST APIs", technical_line)
            self.assertIn("Flask", technical_line)
            self.assertNotIn("Kotlin", technical_line)
            self.assertNotIn("GraphQL", technical_line)

    def test_export_technical_skills_merges_job_matched_skills(self):
        merged = CVDocumentBuilder._export_technical_skills(
            {
                "job_matched_skills": ["REST APIs", "Python"],
                "technical_skills": ["Java", "Python", "Flask"],
            }
        )
        self.assertEqual(merged, ["REST APIs", "Python", "Java", "Flask"])

    def test_build_from_preview_lines_writes_body_and_skips_preview_only_sections(self):
        template_path = get_default_cv_template_path()
        preview_lines = [
            {"text": "Jane Doe", "kind": "name"},
            {"text": "Backend Engineer", "kind": "title"},
            {"text": "Professional Summary", "kind": "section"},
            {"text": "Builds APIs with Python.", "kind": "text"},
            {"text": "Skills Matching Job", "kind": "section"},
            {"text": "• Python • GraphQL", "kind": "skills", "variant": "matched"},
            {"text": "Technical Skills", "kind": "section"},
            {"text": "• Python • Flask", "kind": "skills"},
            {"text": "Experience Highlights", "kind": "section"},
            {"text": "Developer", "kind": "role"},
            {"text": "Acme · 2020-2024", "kind": "meta"},
            {"text": "• Built APIs", "kind": "bullet"},
        ]
        profile = {
            "full_name": "Jane Doe",
            "email": "jane@example.com",
            "phone": "555-0100",
        }

        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = os.path.join(tmpdir, "preview_cv.docx")
            builder = CVDocumentBuilder(template_path)
            builder.build_from_preview_lines(output_path, preview_lines, profile, {})

            from docx import Document

            doc = Document(output_path)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

            self.assertIn("Builds APIs with Python.", paragraphs)
            self.assertIn("Technical Skills", paragraphs)
            self.assertIn("Developer", paragraphs)
            self.assertIn("Built APIs", paragraphs)
            self.assertNotIn("Skills Matching Job", paragraphs)
            self.assertNotIn("GraphQL", " ".join(paragraphs))


if __name__ == "__main__":
    unittest.main()
