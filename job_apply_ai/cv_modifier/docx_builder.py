"""Build a tailored CV document from AI-generated structured content."""

from __future__ import annotations

import logging
import shutil
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

from job_apply_ai.storage.user_profile import parse_professional_titles

logger = logging.getLogger(__name__)

SECTION_KEYWORDS = {
    "summary": [
        "summary",
        "profile",
        "personal summary",
        "professional summary",
        "about me",
        "personal statement",
        "career summary",
    ],
    "skills": [
        "skills",
        "technical skills",
        "core skills",
        "key skills",
        "competencies",
        "expertise",
        "qualifications",
    ],
    "tools": [
        "tool & platforms",
        "tools & platforms",
        "tools and platforms",
        "tool and platforms",
        "tools",
        "platforms",
    ],
    "experience": [
        "experience",
        "work experience",
        "employment history",
        "professional experience",
        "career history",
    ],
    "projects": [
        "personal projects",
        "projects",
        "portfolio",
        "selected projects",
    ],
    "soft_skills": [
        "soft skills",
        "interpersonal skills",
        "personal skills",
    ],
    "languages": [
        "languages",
        "language skills",
    ],
    "education": [
        "education",
        "academic background",
        "certifications",
    ],
}


PREVIEW_ONLY_SECTION_KEYWORDS = [
    "skills matching job description",
    "skills matching job",
    "job-matched skills",
    "matched job skills",
    "skills in job and cv",
    "job skills not in cv",
    "skills not in cv",
    "job requirements not in cv",
    "missing job skills",
]


class CVDocumentBuilder:
    """Apply structured CV content onto a Word template."""

    def __init__(self, template_path: str):
        self.template_path = template_path

    def build(
        self,
        output_path: str,
        content: dict[str, Any],
        profile: dict[str, Any] | None = None,
    ) -> bool:
        shutil.copy2(self.template_path, output_path)
        doc = Document(output_path)

        if profile:
            self._fill_header(doc, profile, content)

        self._remove_preview_only_sections(doc)

        updated = False
        if content.get("professional_summary"):
            updated |= self._replace_section(
                doc,
                SECTION_KEYWORDS["summary"],
                [content["professional_summary"]],
            )

        skills = self._export_technical_skills(content)
        if skills:
            updated |= self._replace_section(
                doc,
                SECTION_KEYWORDS["skills"],
                [self._format_inline_bullets(skills)],
                bullet_style="inline",
            )

        tools = content.get("tools_platforms") or []
        if tools:
            updated |= self._replace_section(
                doc,
                SECTION_KEYWORDS["tools"],
                [self._format_inline_bullets(tools)],
                bullet_style="inline",
            )

        experience_entries = content.get("experience_highlights") or []
        if experience_entries:
            updated |= self._replace_section(
                doc,
                SECTION_KEYWORDS["experience"],
                self._format_experience_entries(experience_entries),
            )

        project_entries = content.get("personal_projects") or []
        if project_entries:
            updated |= self._replace_section(
                doc,
                SECTION_KEYWORDS["projects"],
                self._format_project_entries(project_entries),
            )

        soft_skills = content.get("soft_skills") or []
        if soft_skills:
            updated |= self._replace_section(
                doc,
                SECTION_KEYWORDS["soft_skills"],
                [f"• {skill}" if not str(skill).startswith("•") else str(skill) for skill in soft_skills],
            )

        language_lines = content.get("languages") or []
        if isinstance(language_lines, str):
            language_lines = [language_lines]
        if language_lines:
            updated |= self._replace_section(doc, SECTION_KEYWORDS["languages"], language_lines)

        education_lines = content.get("education") or []
        if isinstance(education_lines, str):
            education_lines = [education_lines]
        if education_lines:
            updated |= self._replace_section(doc, SECTION_KEYWORDS["education"], education_lines)

        additional = content.get("additional_sections") or {}
        for title, lines in additional.items():
            if not lines:
                continue
            if isinstance(lines, str):
                lines = [lines]
            updated |= self._replace_section(doc, [title.lower()], lines)

        if not updated:
            self._append_generated_sections(doc, content)

        doc.save(output_path)
        return True

    def _remove_preview_only_sections(self, doc: Document) -> None:
        """Drop preview-only job skill sections from the exported document."""
        while True:
            start_idx = self._find_section_start(doc, PREVIEW_ONLY_SECTION_KEYWORDS)
            if start_idx is None:
                break
            end_idx = self._find_section_end(doc, start_idx)
            self._remove_paragraph_range(doc, start_idx, end_idx)

    def _fill_header(
        self,
        doc: Document,
        profile: dict[str, Any],
        content: dict[str, Any] | None = None,
    ) -> None:
        if not doc.paragraphs:
            return

        name = profile.get("full_name", "").strip()
        title = str((content or {}).get("professional_title", "")).strip()
        if not title:
            titles = parse_professional_titles(profile.get("professional_title", ""))
            title = titles[0] if len(titles) == 1 else ""
        doc.paragraphs[0].text = f"{name}\t\t{title}" if name and title else name or title

        if len(doc.paragraphs) < 2:
            return

        contact_parts = []
        if profile.get("email"):
            contact_parts.append(f"Email: {profile['email']}")
        if profile.get("github"):
            contact_parts.append(f"GitHub: {profile['github']}")
        if profile.get("phone"):
            contact_parts.append(f"Phone: {profile['phone']}")
        if profile.get("linkedin"):
            contact_parts.append(f"LinkedIn: {profile['linkedin']}")

        if contact_parts:
            doc.paragraphs[1].text = " | ".join(contact_parts)

    @staticmethod
    def _export_technical_skills(content: dict[str, Any]) -> list[str]:
        """Merge job-matched skills into technical skills for exported documents."""
        matched = content.get("job_matched_skills") or []
        technical = content.get("technical_skills") or content.get("key_skills") or []
        merged: list[str] = []
        seen: set[str] = set()
        for skill in [*matched, *technical]:
            text = str(skill).strip()
            if not text:
                continue
            key = text.lower()
            if key in seen:
                continue
            seen.add(key)
            merged.append(text)
        return merged

    @staticmethod
    def _format_inline_bullets(items: list[Any]) -> str:
        """Join skill/tool items on one line with bullet separators."""
        cleaned = [str(item).strip().lstrip("•").strip() for item in items if str(item).strip()]
        if not cleaned:
            return ""
        return " • ".join(f"• {item}" for item in cleaned)

    def _format_experience_entries(self, entries: list[Any]) -> list[str]:
        lines: list[str] = []
        for entry in entries:
            if isinstance(entry, str):
                lines.append(entry)
                continue
            role = entry.get("role") or entry.get("title") or "Role"
            company = entry.get("company") or ""
            period = entry.get("period") or entry.get("dates") or ""
            header_parts = [part for part in [role, company, period] if part]
            lines.append(" | ".join(header_parts))
            for bullet in entry.get("bullets") or entry.get("highlights") or []:
                bullet_text = str(bullet).strip()
                if bullet_text:
                    lines.append(f"• {bullet_text.lstrip('•').strip()}")
            lines.append("")
        return lines

    def _format_project_entries(self, entries: list[Any]) -> list[str]:
        lines: list[str] = []
        for entry in entries:
            if isinstance(entry, str):
                lines.append(entry)
                continue
            name = entry.get("name") or entry.get("title") or "Project"
            description = entry.get("description") or ""
            header_parts = [part for part in [name, description] if part]
            lines.append(" | ".join(header_parts))
            for bullet in entry.get("bullets") or entry.get("highlights") or []:
                bullet_text = str(bullet).strip()
                if bullet_text:
                    lines.append(f"• {bullet_text.lstrip('•').strip()}")
            lines.append("")
        return lines

    def _replace_section(
        self,
        doc: Document,
        keywords: list[str],
        lines: list[str],
        *,
        bullet_style: str = "list",
    ) -> bool:
        start_idx = self._find_section_start(doc, keywords)
        if start_idx is None:
            return False

        end_idx = self._find_section_end(doc, start_idx)
        self._remove_paragraph_range(doc, start_idx + 1, end_idx)

        anchor = doc.paragraphs[start_idx]
        for line in lines:
            if not str(line).strip():
                anchor = self._insert_paragraph_after(anchor, "")
                continue

            text = str(line).strip()
            if bullet_style == "inline":
                anchor = self._insert_paragraph_after(anchor, text)
                anchor.style = "Normal"
                continue

            anchor = self._insert_paragraph_after(anchor, text.lstrip("•").strip())
            if text.startswith("•"):
                anchor.style = "List Bullet"
            elif "|" in text:
                for run in anchor.runs:
                    run.bold = True
        return True

    def _find_section_start(self, doc: Document, keywords: list[str]) -> int | None:
        for index, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.lower().strip()
            if not text:
                continue
            if any(keyword == text or text.startswith(keyword) for keyword in keywords):
                return index
            if paragraph.style.name.startswith("Heading") and any(keyword in text for keyword in keywords):
                return index
        return None

    def _find_section_end(self, doc: Document, start_idx: int) -> int:
        for index in range(start_idx + 1, len(doc.paragraphs)):
            paragraph = doc.paragraphs[index]
            if paragraph.style.name.startswith("Heading") and paragraph.text.strip():
                return index
        return len(doc.paragraphs)

    def _remove_paragraph_range(self, doc: Document, start_idx: int, end_idx: int) -> None:
        for index in reversed(range(start_idx, end_idx)):
            if index < len(doc.paragraphs):
                element = doc.paragraphs[index]._element
                element.getparent().remove(element)

    @staticmethod
    def _insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        new_para = Paragraph(new_p, paragraph._parent)
        if text:
            new_para.add_run(text)
        return new_para

    def _append_generated_sections(self, doc: Document, content: dict[str, Any]) -> None:
        doc.add_page_break()
        heading = doc.add_paragraph("Tailored Application Content")
        heading.style = "Heading 1"

        if content.get("professional_summary"):
            title = doc.add_paragraph("Personal Summary")
            title.style = "Heading 2"
            doc.add_paragraph(content["professional_summary"])

        skills = self._export_technical_skills(content)
        if skills:
            title = doc.add_paragraph("Technical Skills")
            title.style = "Heading 2"
            doc.add_paragraph(self._format_inline_bullets(skills))

        tools = content.get("tools_platforms") or []
        if tools:
            title = doc.add_paragraph("Tools & Platforms")
            title.style = "Heading 2"
            doc.add_paragraph(self._format_inline_bullets(tools))

        experience_entries = content.get("experience_highlights") or []
        if experience_entries:
            title = doc.add_paragraph("Work Experience")
            title.style = "Heading 2"
            for line in self._format_experience_entries(experience_entries):
                if line.startswith("•"):
                    paragraph = doc.add_paragraph(line.lstrip("•").strip())
                    paragraph.style = "List Bullet"
                elif line.strip():
                    paragraph = doc.add_paragraph(line)
                    if "|" in line:
                        for run in paragraph.runs:
                            run.bold = True

        project_entries = content.get("personal_projects") or []
        if project_entries:
            title = doc.add_paragraph("Personal Projects")
            title.style = "Heading 2"
            for line in self._format_project_entries(project_entries):
                if line.startswith("•"):
                    paragraph = doc.add_paragraph(line.lstrip("•").strip())
                    paragraph.style = "List Bullet"
                elif line.strip():
                    paragraph = doc.add_paragraph(line)
                    if "|" in line:
                        for run in paragraph.runs:
                            run.bold = True

        soft_skills = content.get("soft_skills") or []
        if soft_skills:
            title = doc.add_paragraph("Soft Skills")
            title.style = "Heading 2"
            for skill in soft_skills:
                paragraph = doc.add_paragraph(str(skill))
                paragraph.style = "List Bullet"

        language_lines = content.get("languages") or []
        if isinstance(language_lines, str):
            language_lines = [language_lines]
        if language_lines:
            title = doc.add_paragraph("Languages")
            title.style = "Heading 2"
            for line in language_lines:
                doc.add_paragraph(str(line))

        education_lines = content.get("education") or []
        if isinstance(education_lines, str):
            education_lines = [education_lines]
        if education_lines:
            title = doc.add_paragraph("Education")
            title.style = "Heading 2"
            for line in education_lines:
                doc.add_paragraph(str(line))
