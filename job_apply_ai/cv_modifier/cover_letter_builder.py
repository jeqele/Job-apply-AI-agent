"""Build cover letter Word documents from structured content."""

from __future__ import annotations

from typing import Any

from docx import Document
from docx.shared import Pt


class CoverLetterBuilder:
    """Create a cover letter .docx from structured JSON content."""

    def build(self, output_path: str, content: dict[str, Any]) -> None:
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        if content.get("date"):
            doc.add_paragraph(str(content["date"]))

        recipient_lines = []
        if content.get("recipient_name"):
            recipient_lines.append(str(content["recipient_name"]))
        if content.get("recipient_company"):
            recipient_lines.append(str(content["recipient_company"]))
        for line in recipient_lines:
            doc.add_paragraph(line)

        doc.add_paragraph("")
        if content.get("greeting"):
            doc.add_paragraph(str(content["greeting"]))

        for paragraph in content.get("body_paragraphs") or []:
            if str(paragraph).strip():
                doc.add_paragraph(str(paragraph).strip())

        doc.add_paragraph("")
        if content.get("closing"):
            doc.add_paragraph(str(content["closing"]))
        if content.get("signature_name"):
            doc.add_paragraph(str(content["signature_name"]))

        contact_bits = []
        if content.get("candidate_email"):
            contact_bits.append(str(content["candidate_email"]))
        if content.get("candidate_phone"):
            contact_bits.append(str(content["candidate_phone"]))
        if contact_bits:
            doc.add_paragraph(" | ".join(contact_bits))

        doc.save(output_path)
